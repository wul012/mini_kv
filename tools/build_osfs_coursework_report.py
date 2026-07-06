from __future__ import annotations

import textwrap
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DELIVERY = ROOT / "课程设计交付" / "v1633-osfs-final"
IMAGE_DIR = DELIVERY / "images"
REPORT_DOCX = DELIVERY / "操作系统课程设计报告-OSFS-v1633.docx"


def ensure_dirs() -> None:
    DELIVERY.mkdir(parents=True, exist_ok=True)
    IMAGE_DIR.mkdir(parents=True, exist_ok=True)


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/simhei.ttf") if bold else Path("C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def wrap_text(text: str, width: int) -> list[str]:
    lines: list[str] = []
    for paragraph in text.splitlines():
        if not paragraph:
            lines.append("")
            continue
        lines.extend(textwrap.wrap(paragraph, width=width, break_long_words=False, replace_whitespace=False))
    return lines


def draw_box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], text: str, fill: str, outline: str = "#1f2937") -> None:
    draw.rounded_rectangle(xy, radius=16, fill=fill, outline=outline, width=2)
    x1, y1, x2, y2 = xy
    title_font = font(28, bold=True)
    body_font = font(21)
    parts = text.split("\n", 1)
    title = parts[0]
    body = parts[1] if len(parts) > 1 else ""
    draw.text((x1 + 24, y1 + 18), title, fill="#111827", font=title_font)
    if body:
        y = y1 + 58
        for line in wrap_text(body, 22):
            draw.text((x1 + 24, y), line, fill="#374151", font=body_font)
            y += 30


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#374151") -> None:
    draw.line([start, end], fill=color, width=4)
    ex, ey = end
    sx, sy = start
    if ex == sx:
        direction = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 10, ey - 18 * direction), (ex + 10, ey - 18 * direction)]
    else:
        direction = 1 if ex > sx else -1
        points = [(ex, ey), (ex - 18 * direction, ey - 10), (ex - 18 * direction, ey + 10)]
    draw.polygon(points, fill=color)


def create_architecture_diagram() -> Path:
    path = IMAGE_DIR / "01-osfs-system-architecture.png"
    image = Image.new("RGB", (1500, 900), "#f8fafc")
    draw = ImageDraw.Draw(image)
    draw.text((60, 45), "OSFS 课程设计系统结构", fill="#0f172a", font=font(42, bold=True))
    draw.text((60, 100), "命令、文件系统、块设备三层分工；OSFS 与原 KV/WAL/TCP 主线保持隔离。", fill="#475569", font=font(24))

    boxes = [
        ((90, 170, 510, 300), "用户/演示者\n先 LOGIN 认证，再执行 DIR、OPEN、READ、WRITE 等命令。", "#dbeafe"),
        ((90, 365, 510, 515), "CommandProcessor\n维护登录身份与 fd 表；每个句柄保存独立读写位移。", "#e0f2fe"),
        ((90, 590, 510, 750), "FileSystem\n解析 MFD/UFD，执行权限检查、inode/块分配和区间 I/O。", "#dcfce7"),
        ((780, 590, 1200, 750), "VirtualDisk\n把块读写落到二进制镜像文件，模拟磁盘设备。", "#fef3c7"),
        ((780, 365, 1200, 515), "磁盘镜像 osfs-course.img\n持久化用户表、MFD/UFD、双位图、inode 与文件数据。", "#ffedd5"),
        ((780, 170, 1200, 300), "原 mini-kv 主线\nStore、WAL、Snapshot、TCP/RESP 不被 OSFS 牵动。", "#f3e8ff"),
    ]
    for xy, text, fill in boxes:
        draw_box(draw, xy, text, fill)
    draw_arrow(draw, (300, 305), (300, 360))
    draw_arrow(draw, (300, 520), (300, 585))
    draw_arrow(draw, (515, 670), (775, 670))
    draw_arrow(draw, (990, 585), (990, 520))
    draw.line([(580, 140), (580, 790)], fill="#94a3b8", width=3)
    draw.text((605, 790), "边界：OSFS 是课程设计入口，不授予 KV 写路由、WAL、网络或 managed-audit 执行权。", fill="#334155", font=font(22))
    image.save(path)
    return path


def create_disk_layout_diagram() -> Path:
    path = IMAGE_DIR / "02-osfs-disk-layout.png"
    image = Image.new("RGB", (1500, 820), "#ffffff")
    draw = ImageDraw.Draw(image)
    draw.text((60, 45), "虚拟磁盘布局", fill="#111827", font=font(42, bold=True))
    draw.text((60, 100), "二进制文件被切成固定大小块，OSFS 在这些块上组织元数据和文件内容。", fill="#4b5563", font=font(24))

    x, y, h = 80, 190, 120
    widths = [135, 135, 135, 135, 210, 150, 150, 150]
    labels = [
        ("Block 0", "Superblock\n布局与空闲计数"),
        ("Block 1", "Block bitmap\n块占用状态"),
        ("Block 2", "Inode bitmap\ninode 占用状态"),
        ("Block 3", "User table\nuid/gid/hash/UFD"),
        ("Block 4..N", "Inode table\n类型、权限、时间、直接块"),
        ("Data", "MFD\nuser -> UFD inode"),
        ("Data", "UFD blocks\nfilename -> inode"),
        ("Data", "File/free blocks\n正文与空闲空间"),
    ]
    colors = ["#bfdbfe", "#bbf7d0", "#bbf7d0", "#fecaca", "#fde68a", "#fed7aa", "#e0e7ff", "#e9d5ff"]
    cx = x
    for width, (top, bottom), color in zip(widths, labels, colors):
        draw.rounded_rectangle((cx, y, cx + width, y + h), radius=10, fill=color, outline="#1f2937", width=2)
        draw.text((cx + 16, y + 14), top, fill="#111827", font=font(24, bold=True))
        yy = y + 48
        for line in bottom.split("\n"):
            draw.text((cx + 16, yy), line, fill="#374151", font=font(19))
            yy += 26
        cx += width + 14

    detail = [
        ("MFD", "主文件目录保存用户名到 UFD inode 的映射，是二级目录的第一层。"),
        ("UFD", "每个用户拥有独立目录 inode 和数据块，DIR 只读取当前 UFD。"),
        ("User table", "用户、uid/gid、教学密码哈希和 UFD inode 随磁盘镜像持久化。"),
        ("Double bitmap", "块和 inode 分开分配；超级块记录两类空闲计数。"),
    ]
    y2 = 390
    for i, (title, body) in enumerate(detail):
        xx = 100 + (i % 2) * 680
        yy = y2 + (i // 2) * 160
        draw_box(draw, (xx, yy, xx + 600, yy + 115), f"{title}\n{body}", "#f8fafc", "#64748b")
    image.save(path)
    return path


def create_write_flow_diagram() -> Path:
    path = IMAGE_DIR / "03-osfs-descriptor-flow.png"
    image = Image.new("RGB", (1500, 900), "#f8fafc")
    draw = ImageDraw.Draw(image)
    draw.text((60, 45), "文件描述符位移与区间 I/O", fill="#0f172a", font=font(42, bold=True))
    draw.text((60, 100), "fd 保存独立读写位置；底层按块内偏移读取、覆盖或扩展。", fill="#475569", font=font(24))

    nodes = [
        ((80, 190, 390, 300), "1. OPEN\nr / w / rw / a 决定初始位移"),
        ((520, 190, 830, 300), "2. fd 表\nname + readable/writable + offsets"),
        ((960, 190, 1270, 300), "3. READ/WRITE\n校验句柄模式与 inode 权限"),
        ((960, 420, 1270, 540), "4. 块定位\ndirect = offset / blockSize"),
        ((520, 420, 830, 540), "5. 区间 I/O\n跨块读取或覆盖/扩展"),
        ((80, 420, 390, 540), "6. 推进位移\n只按实际成功字节数更新"),
        ((520, 660, 830, 780), "SEEK/TELL\n修改或观察句柄位置；失败写入不推进 offset"),
    ]
    for xy, text in nodes:
        fill = "#fee2e2" if text.startswith("失败") else "#dbeafe"
        draw_box(draw, xy, text, fill)
    draw_arrow(draw, (390, 245), (515, 245))
    draw_arrow(draw, (830, 245), (955, 245))
    draw_arrow(draw, (1115, 305), (1115, 415))
    draw_arrow(draw, (960, 480), (835, 480))
    draw_arrow(draw, (520, 480), (395, 480))
    draw_arrow(draw, (675, 545), (675, 655), "#b91c1c")
    draw.text((900, 610), "关键边界：位移属于句柄，不属于文件 inode。", fill="#991b1b", font=font(24, bold=True))
    image.save(path)
    return path


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_run_font(run, name: str = "SimSun", size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_para(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.first_line_indent = Cm(0.74) if style is None else None
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, "SimSun", 10.5)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    set_run_font(run, "SimHei", 14 if level == 1 else 12, True)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, "SimSun", 9)
    run.font.color.rgb = RGBColor(75, 85, 99)


def add_figure(doc: Document, image_path: Path, caption: str, width_cm: float = 14.0) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Cm(width_cm))
    add_caption(doc, caption)


def read_optional_text(path: Path) -> str:
    if not path.exists():
        return ""
    return path.read_text(encoding="utf-8", errors="replace").strip()


def add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "要求或模块"
    hdr[1].text = "本项目中的对应实现"
    for cell in hdr:
        set_cell_shading(cell, "D9EAF7")
        for para in cell.paragraphs:
            for run in para.runs:
                set_run_font(run, "SimHei", 10.5, True)
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, "SimSun", 10)


def add_page_number(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend([field_begin, instruction, field_end])
    set_run_font(run, "Times New Roman", 9)


def build_report() -> None:
    ensure_dirs()
    architecture = create_architecture_diagram()
    disk_layout = create_disk_layout_diagram()
    descriptor_flow = create_write_flow_diagram()

    demo_image = ROOT / "f" / "1633" / "图片" / "02-final-osfs-demo.png"
    ctest_image = ROOT / "f" / "1633" / "图片" / "03-full-ctest-summary.png"

    transcript = read_optional_text(DELIVERY / "osfs-demo-transcript.txt")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.4)
    add_page_number(section)

    styles = doc.styles
    styles["Normal"].font.name = "SimSun"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)
    run = title.add_run("操作系统课程设计报告")
    set_run_font(run, "SimHei", 24, True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("基于虚拟磁盘的 MFD/UFD 二级文件系统设计与实现")
    set_run_font(run, "SimHei", 18, True)

    info = [
        ("课程名称", "操作系统课程设计"),
        ("实验题目", "实验二 Linux 二级文件系统"),
        ("项目名称", "mini-kv OSFS 二级文件系统"),
        ("学生姓名", "【请填写】"),
        ("学号", "【请填写】"),
        ("班级", "【请填写】"),
        ("指导教师", "【请填写】"),
        ("完成日期", "2026 年 7 月"),
    ]
    table = doc.add_table(rows=len(info), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(info):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        set_cell_shading(table.rows[i].cells[0], "F1F5F9")
        for cell in table.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, "SimSun", 11)

    doc.add_page_break()
    add_heading(doc, "摘要", 1)
    add_para(
        doc,
        "本课程设计实现了一个运行在用户态的二级文件系统 OSFS。程序把普通二进制文件划分为固定大小的"
        "磁盘块，在镜像中保存超级块、block bitmap、inode bitmap、用户表、inode 表、主文件目录 MFD、"
        "用户文件目录 UFD 和文件数据。格式化时预置 root、Alice、Bob 三个用户；用户记录包含 uid、gid、"
        "教学用途的密码哈希和 UFD inode，关闭程序后重新打开镜像仍能完成认证并恢复各自目录。"
    )
    add_para(
        doc,
        "系统提供 LOGIN、DIR、CREATE、DELETE、OPEN、CLOSE、READ、WRITE、CHMOD、STAT、SEEK 和 TELL"
        " 命令。每个 fd 保存独立读写位移，分段读取、连续写入、原位覆盖和追加均由数据块区间 I/O 完成。"
        "测试覆盖错误密码、未登录阻断、用户目录隔离、跨用户权限拒绝、重开持久性、跨块读写、空间不足"
        "保持旧内容和完整命令行演示。OSFS 使用 mini-kv 仓库的构建与测试基础，但不进入 KV、WAL、"
        "snapshot 或 TCP/RESP 执行路径。"
    )
    add_para(doc, "关键词：操作系统；二级目录；MFD/UFD；虚拟磁盘；inode；文件描述符")

    add_heading(doc, "1 课题要求与设计取舍", 1)
    add_para(
        doc,
        "课程资料把实验二定义为“Linux 二级文件系统”：用二进制文件模拟磁盘，保存用户、inode、超级块和"
        "文件信息，完成 dir、create、delete、open、close、read、write、login，并实施读写保护。题目"
        "前一页还提出多用户和多级目录目标。这里的“二级”按经典课程设计含义处理为 MFD/UFD 两级，而不是"
        "只做一个名字叫 root 的平面目录。MFD 记录用户名到 UFD inode 的映射，每个用户的 UFD 再保存本用户"
        "文件名到普通 inode 的映射。"
    )
    add_para(
        doc,
        "mini-kv 原项目已有 WAL 和 snapshot，但这些能力解决的是键值存储恢复，并不能替代文件系统的 inode、"
        "目录和权限。OSFS 因而采用独立 `minikv_osfs` 入口，复用 CMake、CTest 和跨平台构建，不复用 KV"
        " 数据模型。这个取舍让磁盘布局可以按字节解释，也避免为了表面复用而把课程要求藏进键值接口。"
    )

    add_key_value_table(
        doc,
        [
            ("二进制磁盘", "VirtualDisk 按 block number 读写 osfs-course.img，重开后状态仍在。"),
            ("文件系统元数据", "Superblock、block bitmap、inode bitmap、user table、inode table 均有固定落盘区域。"),
            ("二级目录", "MFD 保存 user -> UFD inode；root、Alice、Bob 分别拥有真实 UFD inode 和目录块。"),
            ("用户与 login", "用户表持久化 uid/gid/教学密码哈希/UFD inode；错误密码不会切换身份。"),
            ("文件操作", "DIR、CREATE、DELETE、OPEN、CLOSE、READ、WRITE、CHMOD、STAT 均从真实 CLI 可执行。"),
            ("文件描述符", "Handle 保存 read/write offset；READ/WRITE 推进位置，SEEK/TELL 可修改和观察。"),
            ("读写保护", "权限按 root、owner、group、other 顺序判断，跨用户 0600 文件访问被拒绝。"),
        ],
    )

    add_heading(doc, "2 总体结构设计", 1)
    add_para(
        doc,
        "程序按三层组织。CommandProcessor 解析文本命令，保存登录身份和 fd 表；FileSystem 处理用户目录"
        "解析、权限、inode、位图和区间 I/O；VirtualDisk 只执行定长块读写。磁盘结构的编解码又拆到"
        " osfs_disk_layout 模块，目录和文件内容分别由 osfs_directory、osfs_file_io 管理。这样扩展 fd"
        " 位移时不需要碰用户认证，修改目录时也不会改写块设备接口。"
    )
    add_figure(doc, architecture, "图 1 OSFS 课程设计系统结构", 14.5)

    add_heading(doc, "3 磁盘布局与关键数据结构", 1)
    add_para(
        doc,
        "磁盘格式版本为 2，默认 block size 是 512 字节。Block 0 保存超级块；Block 1 和 Block 2 分别是"
        "块位图、inode 位图；Block 3 保存固定长度用户记录；其后是 inode table，最后才是目录与普通文件"
        "数据。超级块记录各区域起点、总量、空闲块、空闲 inode、用户数和 MFD inode。打开镜像时先校验 magic、"
        "版本和几何信息，v1 镜像会明确要求重新格式化，避免按错误结构解释字节。"
    )
    add_para(
        doc,
        "每个 inode 固定 96 字节，保存类型、mode、owner uid/gid、size、ctime、atime、mtime 和八个直接块号。"
        "目录项固定 64 字节，只保存 used、inode 和 name。用户记录同样为 64 字节，保存用户名、uid、gid、"
        "UFD inode 和密码哈希。定长结构配合 static_assert，能够在编译期发现填充造成的格式漂移。"
    )
    add_figure(doc, disk_layout, "图 2 OSFS 虚拟磁盘布局", 14.5)

    add_heading(doc, "4 格式化、认证与二级目录", 1)
    add_para(
        doc,
        "格式化首先为 MFD 分配 inode 0 和目录块，再为 root、Alice、Bob 各分配一个 UFD inode 与目录块。"
        "MFD 的三个目录项指向这三个 UFD，用户表保存相同 UFD inode。登录时，程序从用户表定位记录，使用"
        "用户名参与教学哈希计算，再比较磁盘中的 hash。认证失败只返回统一错误；认证成功会切换 uid 并清空"
        "旧 fd，防止句柄跨身份遗留。"
    )
    add_para(
        doc,
        "未带用户名前缀的文件名总是在当前用户 UFD 中解析，所以 Alice 创建的 report 不会进入 Bob 的目录。"
        "限定名字 `alice/report` 允许 root 审查，也让测试能够让 Bob 到达目标 inode 后接受权限判断。DIR 默认"
        "只列当前 UFD；普通用户不能列别人的 UFD，root 可以使用 `DIR alice`。这两条规则把目录隔离和 inode"
        "权限分开，避免用简单的输出过滤假装二级目录。"
    )

    add_heading(doc, "5 文件描述符与区间读写", 1)
    add_para(
        doc,
        "fd 不是文件名的另一个写法。Handle 同时保存文件名、可读/可写标志、read offset 和 write offset。"
        "`OPEN w` 先截断；`rw` 保留内容并从位置 0 开始；`a` 把写位置设为 EOF。READ 按实际返回字节数推进"
        "读位置，WRITE 只有成功落盘后才推进写位置。SEEK 设置允许方向的偏移，TELL 输出两个位置。"
    )
    add_para(
        doc,
        "底层区间 I/O 用 `offset / block_size` 选择直接块，用 `offset % block_size` 找块内位置。一段数据跨越"
        "512 字节边界时，循环会在当前块末尾切分，再进入下一个块。扩展写入前先检查八个直接块上限和空闲块，"
        "不足时不修改数据或 fd 位置；新块先清零，避免间隙暴露旧内容。"
    )
    add_figure(doc, descriptor_flow, "图 3 文件描述符位移与区间 I/O", 14.5)

    add_heading(doc, "6 权限模型与失败处理", 1)
    add_para(
        doc,
        "inode mode 使用八进制权限位。判断顺序为 root、owner、group、other；root uid 0 可以审查全部目录，普通"
        "用户只能按自身 uid/gid 和权限位访问。Alice 把文件改成 0600 后，Bob 即使写出 `alice/report` 也会在"
        " inode 权限处得到 permission denied。密码哈希和 mode 都实际参与命令执行，不是 DIR 中的装饰字段。"
    )
    add_para(
        doc,
        "整文件替换先比较“空闲块 + 当前文件可回收块”和目标需求，空间不足时旧内容保持不变。区间扩展写也"
        "在分配前检查所需新增块。创建文件若已取得 inode、随后目录扩容失败，会撤销 inode bitmap；删除则释放"
        "全部数据块和 inode，并重算超级块空闲计数。课设磁盘规模较小，重算比依靠多处分散加减更容易审查。"
    )

    add_heading(doc, "7 关键模块说明", 1)
    add_key_value_table(
        doc,
        [
            ("include/minikv/osfs/virtual_disk.hpp", "声明虚拟磁盘接口，提供 create、open、read_block、write_block。"),
            ("src/osfs_virtual_disk.cpp", "负责磁盘镜像大小校验、块边界校验和二进制读写。"),
            ("src/osfs_disk_layout.*", "固定磁盘结构尺寸，读写 superblock、双位图、用户表、inode 和目录块。"),
            ("src/osfs_filesystem.cpp", "格式化、打开镜像、用户表查询和密码认证。"),
            ("src/osfs_directory.cpp", "MFD/UFD 解析、创建、删除、DIR、STAT、CHMOD 与目录权限。"),
            ("src/osfs_file_io.cpp", "整文件替换、truncate、range read/write、块边界和空间预检。"),
            ("src/osfs_command_processor.cpp", "解析命令，维护登录状态与带 read/write offset 的 fd 表。"),
            ("src/osfs_main.cpp", "提供 minikv_osfs 命令行入口，支持 --disk、--format、--blocks、--script。"),
            ("tests/osfs_tests.cpp", "覆盖认证、二级目录、权限、持久化、offset、跨块 I/O 和失败保持。"),
        ],
    )

    add_heading(doc, "8 测试与演示结果", 1)
    add_para(
        doc,
        "osfs_tests 从 C++ API 检查磁盘格式版本、三个 UFD inode、错误密码、目录隔离、跨用户 0600 拒绝、root"
        " 审查、重新打开后的认证与数据持久性。fd 测试连续写 `hello-` 和 `world`，再分段读取、SEEK 覆盖、"
        "append 追加；另一个 700 字节文件专门验证 512 字节块边界。空间不足测试确认失败重写不会破坏旧内容。"
    )
    if demo_image.exists():
        add_figure(doc, demo_image, "图 4 OSFS 命令行演示结果", 14.5)
    if transcript:
        add_para(
            doc,
            "最终脚本从未登录状态开始，先验证命令阻断和错误密码，再以 Alice 创建、连续写、分段读并 chmod"
            " 0600。Bob 的 DIR 不出现 Alice 文件，限定访问也被拒绝；root 最后列出 Alice UFD 并读取内容。"
            "脚本由真实 minikv_osfs 可执行文件生成 transcript，没有手写成功输出。"
        )
    if ctest_image.exists():
        add_figure(doc, ctest_image, "图 5 全量 CTest 验证证据", 14.5)

    add_heading(doc, "9 运行与复现步骤", 1)
    add_para(
        doc,
        "构建只需要 CMake 和 C++20 编译器。课堂演示启动 OSFS 即可，不需要启动 mini-kv TCP 服务。第一次"
        "运行使用 `--format` 创建 v2 镜像；后续去掉 `--format` 可以验证用户表、UFD 和文件内容确实持久化。"
    )
    code = doc.add_paragraph()
    code.paragraph_format.left_indent = Cm(0.8)
    run = code.add_run(
        "cmake -S . -B build\n"
        "cmake --build build --target minikv_osfs\n"
        "./build/minikv_osfs --disk data/osfs-course.img --format --blocks 96 --script "
        "课程设计交付/v1633-osfs-final/osfs-demo-script.txt"
    )
    set_run_font(run, "Consolas", 9)

    doc.add_page_break()
    add_heading(doc, "10 需求核对与对抗性自审", 1)
    add_key_value_table(
        doc,
        [
            ("保存用户、inode、超级块和文件信息", "完整：四类信息都位于磁盘镜像固定区域，重开测试通过。"),
            ("二级目录", "完整：MFD 指向 root/Alice/Bob 的真实 UFD inode 与目录块。"),
            ("dir/create/delete/open/close/read/write", "完整：真实 CLI 与 focused CTest 覆盖。"),
            ("login 用户登录", "完整：磁盘用户表与密码校验；错误密码失败且不切换身份。"),
            ("DIR 展示文件名、物理地址、保护码、长度", "完整：同时展示 inode、owner 与 size。"),
            ("读写保护", "完整：chmod 0600 后跨用户读写拒绝。"),
            ("EXT2 完整组与间接块", "部分/教学简化：单 block group、八个直接块，不声称完整 EXT2。"),
        ],
    )
    add_para(
        doc,
        "评分老师最可能提出的质疑是：“这是否只是按 owner 过滤一个根目录，并没有真正的二级目录？”回应不靠"
        "演示文字：格式化会为 MFD 与三个 UFD 分配不同 inode 和数据块，用户表与 MFD 都保存 UFD inode；Alice"
        " 和 Bob 的目录项写入不同 UFD。测试断言三个 UFD inode 不同、双方 DIR 相互隔离，并在重开镜像后再次"
        "认证和读取。root 的 `DIR alice` 又能从 MFD/用户表定位 Alice UFD，说明目录关系存在于磁盘结构。"
    )

    add_heading(doc, "11 设计边界与后续工作", 1)
    add_para(
        doc,
        "OSFS 不是完整 EXT2。它只实现课程要求的 MFD/UFD 两级，不支持任意深度子目录；inode 只有八个直接块，"
        "没有间接块；用户表容量固定，FNV-1a 哈希仅用于教学；系统没有 fsck、崩溃日志和多进程并发。这些限制"
        "不会被包装成已完成项。若继续扩展，优先增加 fsck 检查位图、inode 与目录项一致性，再考虑间接块和"
        "更安全的密码派生。"
    )

    add_heading(doc, "12 总结", 1)
    add_para(
        doc,
        "本次实现把文件拆成了几条可以追踪的关系：用户表把身份连到 UFD，目录项把名字连到 inode，inode 再把"
        "权限、长度和数据块连起来。LOGIN 决定从哪个 UFD 开始查找，OPEN 产生带位置的 fd，READ/WRITE 最终"
        "落到具体块内偏移。沿着这条路径，课程里的超级块、位图、inode、目录、权限和打开文件表不再是分散概念。"
    )
    add_para(
        doc,
        "项目最后保留了明确范围：只做 MFD/UFD 二级目录和直接块，不借完整 EXT2 名义夸大成果。OSFS 与原 KV"
        " 主线仍然隔离，报告中的每个完成项都能指向源码、测试或真实 transcript。"
    )

    add_heading(doc, "参考资料", 1)
    refs = [
        "[1] 操作系统课程设计要求 2026：实验二 Linux 二级文件系统，课程讲义 PDF。",
        "[2] mini-kv OSFS 源码：include/minikv/osfs、src/osfs_*.cpp。",
        "[3] mini-kv OSFS 测试与演示：tests/osfs_tests.cpp、tests/osfs_smoke_script.txt。",
    ]
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("\n".join(refs))
    set_run_font(run, "SimSun", 9)

    doc.save(REPORT_DOCX)


if __name__ == "__main__":
    build_report()
    print(REPORT_DOCX)
